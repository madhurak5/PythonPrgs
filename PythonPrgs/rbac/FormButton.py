# import tkinter as tk
# from openpyxl import load_workbook
#
# def create_buttons(data_column, form_label):
#     """
#     Creates buttons for each value in the specified Excel column and assigns click actions.
#
#     Args:
#         data_column (list): A list of data values from the Excel column.
#         form_label (tk.Label): The label widget in the form where the button text will be inserted.
#     """
#
#     for value in data_column:
#         button = tk.Button(window, text=value, command=lambda val=value: update_form(val, form_label))
#         button.pack(padx=5, pady=5)
#
# def update_form(button_text, form_label):
#     """
#     Updates the form label with the text from the clicked button.
#
#     Args:
#         button_text (str): The text from the clicked button.
#         form_label (tk.Label): The label widget in the form.
#     """
#
#     form_label.config(text=button_text)
#
# # def main():
#     # Replace with your actual Excel file path
# excel_file = "C:/3Bayode/Labs/Consolidated Labs.xlsx"
# sheet_name = "oetexport_aus"  # Replace with the sheet containing your data
#
# # Load data from Excel
# wb = load_workbook(filename=excel_file)
# sheet = wb[sheet_name]
# data_column = [cell.value for cell in sheet["A"]]  # Assuming data is in column A
#
# window = tk.Tk()
# window.title("Button Generator")
#
# # Create a form label
# form_label = tk.Label(window, text="Selected Value:")
# form_label.pack(pady=10)
#
# # Generate buttons and assign click actions
# create_buttons(data_column, form_label)
#
# window.mainloop()
#
import datetime
import tkinter as tk
import docx
from docx import Document
from docx.text.paragraph import Paragraph

from docx.table import Table
import smtplib
#
def create_button(button_name):
    """
    Creates a button with the given name and assigns a click action.

    Args:
        button_name (str): The text to display on the button.
    """

    def replace_and_send():
        """
        Replaces "<firmname>" in the Word document and sends the email.
        """

        # Replace the placeholder with the button text
        doc = Document("C:/3Bayode/Service Contract.docx")  # Replace with your template path
        # today = datetime.datetime.today()
        # formatted_date = today.strftime("%B %d, %Y")
        # print(formatted_date)
        for paragraph in doc.paragraphs:
            if "<firmname>" in paragraph.text:
                paragraph.text = paragraph.text.replace("<firmname>", button_name)
            if "<todaydate>" in paragraph.text:
                today = datetime.datetime.today()
                formatted_date = today.strftime("%B %d, %Y")
                print(formatted_date)
                paragraph.text = paragraph.text.replace("<todaydate>", formatted_date)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "firmname" in paragraph.text:
                            paragraph.text = paragraph.text.replace("<firmname>", button_name)
                        if "<todaydate>" in paragraph.text:
                            today = datetime.datetime.today()
                            formatted_date = today.strftime("%B %d, %Y")
                            print(formatted_date)
                            paragraph.text = paragraph.text.replace("<todaydate>", formatted_date)



        # Save the modified document
        doc.save("modified_document5.docx")  # Replace with desired output path

        # Send the email
        # sender_email = ""  # Replace with your email address
        # sender_password = ""  # Replace with your email password (consider using app passwords for security)
        # receiver_email = ""
        # subject = "Modified Document with Firm Name"
        # message = f"Hi,\nThe attached document has been modified with the firm name: {button_name}.\n"

        # with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        #     server.login(sender_email, sender_password)
        #     server.sendmail(sender_email, receiver_email, f"Subject: {subject}\n\n{message}".encode('utf-8'), filename="modified_document.docx")  # Attach the document

        # print("Email sent successfully!")

    window.title("Firm Name Button")
    button = tk.Button(window, text=button_name, command=replace_and_send)
    button.pack()

# if __name__ == "__main__":
button_name = "Matrox Conformity Group"  # Replace with the desired button text
window = tk.Tk()
create_button(button_name)
window.mainloop()
